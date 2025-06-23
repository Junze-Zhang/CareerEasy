from CareerEasy.llm_utils import llm_request
from CareerEasyBackend.models import Candidate, JobPosting
from CareerEasy.constants import *
from django.conf import settings
from django.core.management import call_command
import django
from datetime import datetime
import json
import os
import re
from typing import Union, Dict, Tuple, List, Any
import numpy as np
from openai import OpenAI
from django.db.models import QuerySet
from presidio_analyzer import AnalyzerEngine
from presidio_anonymizer import AnonymizerEngine
from presidio_anonymizer.entities import OperatorConfig
from presidio_analyzer.nlp_engine import NlpEngineProvider
import threading
from concurrent.futures import ThreadPoolExecutor
import logging
import boto3
from botocore.exceptions import NoCredentialsError, ClientError
import uuid
from urllib.parse import urlparse
import PyPDF2
import docx
from io import BytesIO
import traceback

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "CareerEasy.settings")

django.setup()

# Configure the NLP engine with the medium model
nlp_configuration = {
    "nlp_engine_name": "spacy",
    "models": [{"lang_code": "en", "model_name": "en_core_web_md"}]
}

# Create a thread pool for anonymization tasks
MAX_WORKERS = 3  # Adjust this based on your server's capacity
anonymizer_thread_pool = ThreadPoolExecutor(max_workers=MAX_WORKERS)

print("Loading NLP engine for Presidio anonymizer...")
# Create singleton instances
nlp_engine_provider = NlpEngineProvider(nlp_configuration=nlp_configuration)
nlp_engine = nlp_engine_provider.create_engine()
analyzer = AnalyzerEngine(nlp_engine=nlp_engine)
anonymizer = AnonymizerEngine()
anonymizer_lock = threading.Lock()
print("NLP engine loaded.")


def _find_potential_json(response: str) -> List[str]:
    # Regex pattern to match non-greedily
    pattern = r"\{.*?\}"

    # Find all matches
    matches = re.findall(pattern, response)

    return matches


def validate_exp(response: str) -> bool:
    all_matches = [response] + _find_potential_json(response)
    for match in all_matches:
        try:
            parsed = json.loads(match)
            if "error" in parsed:
                return True
            if isinstance(parsed["experience_months"], int) \
                    and parsed["highest_education"].lower() in ["bachelor",
                                                                "master",
                                                                "phd",
                                                                "high school",
                                                                "below high school",
                                                                "other",
                                                                "bachelor's",
                                                                "master's",
                                                                "doctorate",
                                                                "doctor's"]:
                return True
        except:
            continue
    return False


def validate_skills(response: str) -> bool:
    all_matches = [response] + _find_potential_json(response)
    for match in all_matches:
        try:
            parsed = json.loads(match)
            if isinstance(parsed["skills"], list):
                return True
        except:
            continue
    return False


def validate_ai_highlights(response: str) -> bool:
    all_matches = [response] + _find_potential_json(response)
    for match in all_matches:
        try:
            parsed = json.loads(match)
            if "Error" in parsed or \
                    (isinstance(parsed["highlights"], list) and len(parsed["highlights"]) == NUM_AI_HIGHLIGHTS):
                return True
        except:
            continue
    return False


def extract_from_resume(candidate: Candidate) -> tuple[dict[str, Any], list[str]]:
    llm_client = OpenAI(api_key=os.getenv(
        "DEEPSEEK_API_KEY"), base_url=DEEPSEEK_API_URL)
    candidate_info = {
        "exp_month": None,
        "skills": None,
        "ai_highlights": None,
        "highest_education": None,
    }
    today_month = datetime.now().strftime("%B %Y")
    prompt_exp = PROMPT_CANDIDATE_EXP.format(resume=candidate.anonymous_resume,
                                             today_month=today_month)
    messages = [{"role": "user", "content": prompt_exp}]
    response = llm_request(llm_client,
                           messages=messages,
                           validate_fn=validate_exp,
                           model="deepseek-reasoner",
                           return_raw=True)
    if not response:
        raise ValueError("Failed to extract experience from resume.")
    if "error" in response:
        raise ValueError(json.loads(response)["error"])
    if settings.DEBUG:
        debug_dir = os.path.join(settings.BASE_DIR, '.debug')
        if not os.path.exists(debug_dir):
            os.makedirs(debug_dir)
        with open(os.path.join(debug_dir, f"{candidate.id}-{uuid.uuid4().hex}.txt"), "w") as f:
            f.write(f"Prompt: {prompt_exp}\n")
            # f.write(f"Reasoning: {response.reasoning_content}\n")
            f.write(f"Response: {response.content}\n\n")
        print(response.content)

    all_matches = [response.content] + _find_potential_json(response.content)
    errors = []
    for match in all_matches:
        try:
            parsed = json.loads(match)
            if "Error" in parsed:
                errors.append(parsed["error"])
            candidate_info["exp_month"] = parsed["experience_months"]
            candidate_info["highest_education"] = parsed["highest_education"]
        except Exception as e:
            errors.append(repr(e))

    messages.append({"role": "assistant", "content": response.content})
    messages.append({"role": "user", "content": PROMPT_CANDIDATE_SKILLS})
    response = llm_request(llm_client,
                           messages=messages,
                           validate_fn=validate_skills,
                           model="deepseek-reasoner",
                           return_raw=True)
    if not response:
        raise ValueError("Failed to extract skills from resume.")
    if settings.DEBUG:
        debug_dir = os.path.join(settings.BASE_DIR, '.debug')
        if not os.path.exists(debug_dir):
            os.makedirs(debug_dir)
        with open(os.path.join(debug_dir, f"{candidate.id}-{uuid.uuid4().hex}.txt"), "a") as f:
            f.write(f"Prompt: {PROMPT_CANDIDATE_SKILLS}\n")
            f.write(f"Reasoning: {response.reasoning_content}\n")
            f.write(f"Response: {response.content}\n\n")
        print(response.content)

    all_matches = [response.content] + _find_potential_json(response.content)
    for match in all_matches:
        try:
            parsed = json.loads(match)
            if "Error" in parsed:
                errors.append(parsed["error"])
            candidate_info["skills"] = parsed["skills"]
        except Exception as e:
            errors.append(repr(e))

    messages.append({"role": "assistant", "content": response.content})
    prompt_ai_highlights = PROMPT_CANDIDATE_AI_HIGHLIGHTS.format(
        num_highlights=NUM_AI_HIGHLIGHTS)
    messages.append({"role": "user", "content": prompt_ai_highlights})
    response = llm_request(llm_client,
                           messages=messages,
                           validate_fn=validate_ai_highlights,
                           model="deepseek-reasoner",
                           return_raw=True)
    if not response:
        raise ValueError("Failed to extract highlights from resume.")
    if settings.DEBUG:
        debug_dir = os.path.join(settings.BASE_DIR, '.debug')
        if not os.path.exists(debug_dir):
            os.makedirs(debug_dir)
        with open(os.path.join(debug_dir, f"{candidate.id}-{uuid.uuid4().hex}.txt"), "a") as f:
            f.write(f"Prompt: {prompt_ai_highlights}\n")
            f.write(f"Reasoning: {response.reasoning_content}\n")
            f.write(f"Response: {response.content}\n\n")
        print(response.content)

    all_matches = [response.content] + _find_potential_json(response.content)
    for match in all_matches:
        try:
            parsed = json.loads(match)
            if "Error" in parsed:
                errors.append(parsed["error"])
            candidate_info["ai_highlights"] = parsed["highlights"]
        except Exception as e:
            errors.append(repr(e))

    return candidate_info, errors


def update_ai_highlights(candidate: Candidate, custom_prompt: str) -> dict:
    llm_client = OpenAI(api_key=os.getenv(
        "DEEPSEEK_API_KEY"), base_url=DEEPSEEK_API_URL)
    prompt = PROMPT_CANDIDATE_AI_HIGHLIGHTS_CUSTOMIZE.format(
        num_highlights=NUM_AI_HIGHLIGHTS,
        resume=candidate.anonymous_resume,
        candidate_prompt=custom_prompt
    )
    if settings.DEBUG:
        print(prompt)
    messages = [{"role": "user", "content": prompt}]
    response = llm_request(llm_client,
                           messages=messages,
                           validate_fn=validate_ai_highlights,
                           model="deepseek-reasoner")
    if not response:
        raise ValueError("Failed to update ai highlights. Please try again later.")
    response = json.loads(response)
    if "error" in response:
        raise ValueError(response["error"])
    return response["highlights"]


def validate_query(response: str) -> bool:
    try:
        response = json.loads(response)
        if "error" in response:
            return True
        return isinstance(response["minimal_years_of_experience"], int) and isinstance(
            response["maximal_years_of_experience"], int) and isinstance(response["preferred_title_keywords"],
                                                                         list) and isinstance(
            response["high_priority_keywords"], list) and isinstance(response["low_priority_keywords"], list)
    except:
        return False


def natural_language_to_query(natural_language_query: str) -> dict:
    """
    Convert a natural language query to a SQL query using an LLM.
    Returns a dict with the following keys:
    - minimal_years_of_experience: int
    - maximal_years_of_experience: int
    - preferred_title_keywords: list
    - high_priority_keywords: list
    - low_priority_keywords: list
    """
    prompt = PROMPT_NLQ.format(query=natural_language_query)
    llm_client = OpenAI(api_key=os.getenv(
        "DEEPSEEK_API_KEY"), base_url=DEEPSEEK_API_URL)
    messages = [{"role": "user", "content": prompt}]
    response = llm_request(llm_client,
                           messages=messages,
                           validate_fn=validate_query)
    if not response:
        raise ValueError("Search failed.")
    if "error" in response:
        raise ValueError(json.loads(response)["error"])
    response_json = json.loads(response)
    response_json["preferred_title_keywords"] = [re.sub(
        r'[^a-zA-Z0-9\s]', '', keyword).lower() for keyword in response_json["preferred_title_keywords"]]
    response_json["high_priority_keywords"] = [re.sub(
        r'[^a-zA-Z0-9\s]', '', keyword).lower() for keyword in response_json["high_priority_keywords"]]
    response_json["low_priority_keywords"] = [re.sub(
        r'[^a-zA-Z0-9\s]', '', keyword).lower() for keyword in response_json["low_priority_keywords"]]
    return response_json


# def _match_experience(query: dict, candidates) -> list:
#     min_months = query.get('minimal_years_of_experience', 0) * 12
#     max_months = query.get('maximal_years_of_experience', 100) * 12
#     candidates = candidates.annotate(
#         # Base score: every year of experience is worth 1 point, max 10
#         experience_base_score=Case(
#             When(
#                 experience_months__gte=120,  # 10 years
#                 then=Value(10)
#             ),
#             default=F('experience_months') /
#             Value(12),  # Convert months to years
#             output_field=FloatField()
#         ),
#         # Experience match score: as follows, max 100
#         experience_match_score=Case(
#             # Perfect match (within 3 months)
#             When(
#                 experience_months__gte=min_months - 3,
#                 experience_months__lte=max_months + 3,
#                 then=Value(100)
#             ),
#             # Overqualified; decrease 50 points if overqualified by 2 times (else linearly decrease up to 50);
#             # also decrease 10 points per year overqualified
#             When(
#                 experience_months__gt=max_months,
#                 then=ExpressionWrapper(
#                     Value(50) * Case(
#                         When(
#                             experience_months__gt=Value(max_months * 2),
#                             then=Value(0)
#                         ),
#                         default=Value(2) - F('experience_months') /
#                         Value(max_months),
#                         output_field=FloatField()
#                     ) + Value(50) - Value(10) * (F('experience_months') - Value(max_months)) / Value(12),
#                     output_field=FloatField()
#                 )
#             ),
#             # Underqualified; scale linearly between 0 and 80
#             When(
#                 experience_months__lt=min_months,
#                 then=ExpressionWrapper(
#                     Value(80) - Value(80) * (Value(min_months) -
#                                              F('experience_months')) / Value(min_months),
#                     output_field=FloatField()
#                 )
#             ),
#             # Default case: only possible if min_months>max_months
#             default=Value(100),
#             output_field=FloatField()
#         )
#     ).annotate(
#         # Final experience score = base score + match score * 0.9
#         experience_score=ExpressionWrapper(
#             F('experience_base_score') + Value(0.9) *
#             F('experience_match_score'),
#             output_field=FloatField()
#         )
#     )
#     return candidates


# def _match_title(query: dict, candidates) -> list:
#     query = SearchQuery(query['preferred_title_keywords'])
#     candidates = candidates.annotate(
#         title_match_score=SearchRank(F('title'), query, )
#     )
#     return candidates

# def _match_candidate_info(query: dict, candidates) -> list:
#     query_high = SearchQuery(query['high_priority_keywords'])
#     query_low = SearchQuery(query['low_priority_keywords'])
#     search_vector = SearchVector('title', 'skills', 'ai_highlights', 'highest_education')
#     candidates = candidates.annotate(
#         high_priority_match_score=SearchRank(F('title'), query_high),
#         low_priority_match_score=SearchRank(F('title'), query_low)
#     )
#     return candidates


def _match_score(query: dict, candidate: Dict, return_raw: bool = False) -> Union[Dict, float]:
    experience_score = 0
    title_score = 0
    skills_score = [0, 0]
    ai_highlights_score = [0, 0]
    resume_score = [0, 0]

    candidate_exp_months = candidate['experience_months']
    min_months = query.get('minimal_years_of_experience', 0) * 12
    max_months = query.get('maximal_years_of_experience', 100) * 12
    if min_months > max_months:
        min_months, max_months = max_months, min_months
    if min_months - 3 <= candidate_exp_months <= max_months + 3:  # Perfect match; 3-month margin of error
        experience_score = 100
    elif candidate_exp_months > max_months:  # overqualified
        if max_months == 0:
            experience_score = 0
        else:
            # deduct 50 points if candidate has 2x more than max required (else linearly);
            # also deduct 10 points for each year exceeded (max 50)
            experience_score = max((50 * (2 - candidate_exp_months / max_months)), 0) + \
                               max(50 - 10 * (candidate_exp_months - max_months) / 12, 0)
    else:  # underqualified; scale experience score linearly up to 80
        experience_score = (min_months -
                            candidate_exp_months) / min_months * 80
    # 90% weight on experience match score
    # 10% weight on base score: 10 point for every year of experience, max 100
    experience_score = experience_score * 0.9 + \
                       min(candidate_exp_months / 12 * 10, 100) * 0.1

    # Title matching score: if any of preferred title matches the candidate, they get a score of 100
    query_preferred_title_keywords = query.get('preferred_title_keywords', [])
    for keyword in query_preferred_title_keywords:
        if keyword in candidate['standardized_title']:
            title_score = 100
            break

    # Keyword-matching scores are the recall percentages of every preferred keyword in candidate data
    query_high_priority_keywords = query.get('high_priority_keywords', [])
    query_low_priority_keywords = query.get('low_priority_keywords', [])
    if candidate['standardized_skills']:
        for keyword in query_high_priority_keywords:
            if keyword in candidate['standardized_skills']:
                skills_score[0] += 1
        if candidate['standardized_ai_highlights']:
            for highlight in candidate['standardized_ai_highlights']:
                if keyword in highlight:
                    ai_highlights_score[0] += 1
        if candidate['standardized_anonymous_resume']:
            if keyword in candidate['standardized_anonymous_resume']:
                resume_score[0] += 1
    skills_score[0] = skills_score[0] / \
                      len(query_high_priority_keywords) * \
                      100 if query_high_priority_keywords else 0
    ai_highlights_score[0] = ai_highlights_score[0] / \
                             len(query_high_priority_keywords) * \
                             100 if query_high_priority_keywords else 0
    resume_score[0] = resume_score[0] / \
                      len(query_high_priority_keywords) * \
                      100 if query_high_priority_keywords else 0

    for keyword in query_low_priority_keywords:
        if candidate['standardized_skills']:
            if keyword in candidate['standardized_skills']:
                skills_score[1] += 1
        if candidate['standardized_ai_highlights']:
            for highlight in candidate['standardized_ai_highlights']:
                if keyword in highlight:
                    ai_highlights_score[1] += 1
        if candidate['standardized_anonymous_resume']:
            if keyword in candidate['standardized_anonymous_resume']:
                resume_score[1] += 1
    skills_score[1] = skills_score[1] / \
                      len(query_low_priority_keywords) * \
                      100 if query_low_priority_keywords else 0
    ai_highlights_score[1] = ai_highlights_score[1] / \
                             len(query_low_priority_keywords) * \
                             100 if query_low_priority_keywords else 0
    resume_score[1] = resume_score[1] / \
                      len(query_low_priority_keywords) * \
                      100 if query_low_priority_keywords else 0

    if return_raw:
        return {
            "experience_score": experience_score,
            "title_score": title_score,
            "skills_score": skills_score,
            "ai_highlights_score": ai_highlights_score,
            "resume_score": resume_score
        }
    return float(np.average([experience_score,
                             title_score,
                             *skills_score,
                             *ai_highlights_score,
                             *resume_score], weights=[
        10,
        10,
        5,
        2,
        5,
        2,
        2.5,
        1
    ]))


def rank_candidates(query: dict, candidates: list[Dict]) -> list[Tuple[Dict, float]]:
    candidates_with_score = [{**candidate,
                              "id": str(candidate['id']),
                              "name": candidate['first_name'] + " " + (
                                  candidate['middle_name'][0] + ". " if candidate['middle_name'] else "") + candidate[
                                          'last_name'],
                              "match_score": _match_score(query, candidate)}
                             for candidate in candidates]
    return sorted(candidates_with_score, key=lambda x: x['match_score'], reverse=True)


def am_i_a_match(candidate: Candidate, job: JobPosting) -> bool:
    llm_client = OpenAI(api_key=os.getenv(
        "DEEPSEEK_API_KEY"), base_url=DEEPSEEK_API_URL)
    prompt = PROMPT_CHECK_FIT.format(
        company=job.company.name,
        category=job.company.category.name,
        job_title=job.title,
        job_description=job.description,
        candidate_title=candidate.title,
        highlights="\n".join(candidate.ai_highlights),
        resume=candidate.anonymous_resume
    )
    if settings.DEBUG:
        print(prompt)
    response = llm_request(llm_client,
                           messages=[{"role": "user", "content": prompt}],
                           validate_fn=lambda _: True)
    if not response:
        raise ValueError("Request failed, please try again later.")
    if "error" in response:
        raise ValueError(json.loads(response)["error"])
    return response


def anonymize_resume(resume: str) -> str:
    try:
        resume_lower = resume.lower()
        if "skills" in resume_lower:
            skills_idx = resume_lower.index("skills")
        else:
            skills_idx = 2147483647
        if "experience" in resume_lower:
            experience_idx = resume_lower.index("experience")
        else:
            experience_idx = 2147483647
        if "project" in resume_lower:
            project_idx = resume_lower.index("project")
        else:
            project_idx = 2147483647
        if "education" in resume_lower:
            education_idx = resume_lower.index("education")
        else:
            education_idx = 2147483647
        personal_info_idx = min(skills_idx, experience_idx, project_idx, education_idx)
        resume_personal_info = resume[:personal_info_idx]
        resume_other = resume[personal_info_idx:]

        # Use the thread pool for analysis
        def analyze_text():
            return analyzer.analyze(
                text=resume_personal_info.lower(),
                entities=[
                    "PERSON",
                    "EMAIL_ADDRESS",
                    "PHONE_NUMBER",
                    "SSN",
                    "CREDIT_CARD",
                    "URL",
                    "LOCATION"
                ],
                language='en',
                allow_list=open("technical_terms.txt", "r").read().split(",")
            )

        # Submit the analysis task to the thread pool
        future = anonymizer_thread_pool.submit(analyze_text)
        results = future.result(timeout=30)  # Add timeout to prevent hanging

        with anonymizer_lock:
            anonymized_personal_info = anonymizer.anonymize(
                text=resume_personal_info,
                analyzer_results=results,
                operators={
                    "PHONE_NUMBER": OperatorConfig(
                        "replace",
                        {"new_value": "XXX-XXX-XXXX"}
                    ),
                    "EMAIL_ADDRESS": OperatorConfig(
                        "replace",
                        {"new_value": "<email>"}
                    ),
                    "PERSON": OperatorConfig(
                        "replace",
                        {"new_value": "<name>"}
                    ),
                    "URL": OperatorConfig(
                        "replace",
                        {"new_value": "<url>"}
                    ),
                    "LOCATION": OperatorConfig(
                        "replace",
                        {"new_value": "Anytown, Anycountry"}
                    ),
                    "CREDIT_CARD": OperatorConfig(
                        "replace",
                        {"new_value": "XXXX-XXXX-XXXX-XXXX"}
                    ),
                    "SSN": OperatorConfig(
                        "replace",
                        {"new_value": "XXX-XX-XXXX"}
                    )
                }
            )

        return anonymized_personal_info.text + resume_other
    except Exception as e:
        logger.error(f"Error in anonymize_resume: {str(e)}")
        # Return the original resume if anonymization fails
        return resume


# Add cleanup function to be called when the application shuts down
def cleanup():
    anonymizer_thread_pool.shutdown(wait=True)
    logger.info("Anonymizer thread pool shut down")


def extract_text_from_pdf(file) -> str:
    """
    Extract text from PDF file.
    
    Args:
        file: Django uploaded file object
        
    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: If PDF cannot be processed
    """
    try:
        # Create a BytesIO object from the file
        file_content = BytesIO(file.read())

        # Create a PDF reader object
        pdf_reader = PyPDF2.PdfReader(file_content)

        # Extract text from all pages
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"

        # Reset file pointer for potential future use
        file.seek(0)

        if not text.strip():
            raise ValueError("No text content found in PDF")

        logger.info("Successfully extracted text from PDF")
        return text.strip()

    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file) -> str:
    """
    Extract text from DOCX file.
    
    Args:
        file: Django uploaded file object
        
    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: If DOCX cannot be processed
    """
    try:
        # Create a BytesIO object from the file
        file_content = BytesIO(file.read())

        # Create a document object
        doc = docx.Document(file_content)

        # Extract text from all paragraphs
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"

        # Reset file pointer for potential future use
        file.seek(0)

        if not text.strip():
            raise ValueError("No text content found in DOCX")

        logger.info("Successfully extracted text from DOCX")
        return text.strip()

    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_doc(file) -> str:
    """
    Extract text from DOC file using python-docx2txt.
    Note: DOC files are legacy format and extraction may be limited.
    
    Args:
        file: Django uploaded file object
        
    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: If DOC cannot be processed
    """
    try:
        import docx2txt

        # Create a temporary file to work with docx2txt
        import tempfile

        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
            # Write file content to temporary file
            for chunk in file.chunks():
                temp_file.write(chunk)
            temp_file.flush()

            # Extract text using docx2txt
            text = docx2txt.process(temp_file.name)

        # Clean up temporary file
        os.unlink(temp_file.name)

        # Reset file pointer for potential future use
        file.seek(0)

        if not text.strip():
            raise ValueError("No text content found in DOC")

        logger.info("Successfully extracted text from DOC")
        return text.strip()

    except ImportError:
        logger.error("docx2txt library not available for DOC files")
        raise ValueError("DOC file format not supported. Please convert to DOCX or PDF.")
    except Exception as e:
        logger.error(f"Error extracting text from DOC: {str(e)}")
        raise ValueError(f"Failed to extract text from DOC: {str(e)}")


def extract_text_from_file(file) -> str:
    """
    Extract text from various file formats (TXT, PDF, DOCX, DOC).
    
    Args:
        file: Django uploaded file object
        
    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: If file format is not supported or cannot be processed
    """
    if not hasattr(file, 'name') or not file.name:
        raise ValueError("File must have a name/extension")

    file_extension = os.path.splitext(file.name)[1].lower()

    try:
        if file_extension == '.txt':
            # Handle plain text files
            return file.read().decode('utf-8')
        if file_extension == '.pdf':
            return extract_text_from_pdf(file)
        elif file_extension == '.docx':
            return extract_text_from_docx(file)
        elif file_extension == '.doc':
            return extract_text_from_doc(file)
        else:
            # Defaults to plain text files
            return file.read().decode('utf-8')

    except UnicodeDecodeError:
        raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: .pdf, .docx, .doc, plain text")


def upload_profile_picture_to_s3(file, candidate_id: str) -> str:
    """
    Upload a profile picture to S3 and return the URL.
    
    Args:
        file: Django uploaded file object
        candidate_id: String ID of the candidate
        
    Returns:
        str: S3 URL of the uploaded file
        
    Raises:
        ValueError: If upload fails or invalid file type
    """
    # Validate file type
    allowed_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.webp']
    file_extension = None

    if hasattr(file, 'name') and file.name:
        file_extension = os.path.splitext(file.name)[1].lower()
        if file_extension not in allowed_extensions:
            raise ValueError(f"Invalid file type. Allowed types: {', '.join(allowed_extensions)}")
    else:
        # Default to .jpg if no extension is provided
        file_extension = '.jpg'

    # Validate file size (max 10MB)
    max_size = 10 * 1024 * 1024  # 10MB
    if hasattr(file, 'size') and file.size > max_size:
        raise ValueError("File size too large. Maximum size is 10MB.")

    try:
        # Initialize S3 client
        s3_client = boto3.client(
            's3',
            aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
            aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY'),
            region_name=os.getenv('AWS_REGION', 'ca-central-1')
        )

        # Generate unique filename
        unique_filename = f"profile-pictures/{candidate_id}_{uuid.uuid4().hex[:8]}{file_extension}"

        # Get bucket name from S3_BASE_URL
        from CareerEasyBackend.initDB import S3_BASE_URL
        parsed_url = urlparse(S3_BASE_URL.format(n="dummy"))
        bucket_name = parsed_url.netloc.split('.')[0]  # Extract bucket name from hostname

        # Upload file to S3
        s3_client.upload_fileobj(
            file,
            bucket_name,
            unique_filename,
            ExtraArgs={
                'ContentType': f'image/{file_extension.lstrip(".")}',
                'ACL': 'public-read'  # Make the file publicly readable
            }
        )

        # Return the full S3 URL
        s3_url = f"https://{bucket_name}.s3.ca-central-1.amazonaws.com/{unique_filename}"
        logger.info(f"Successfully uploaded profile picture to S3: {s3_url}")
        return s3_url

    except NoCredentialsError:
        logger.error("AWS credentials not found")
        raise ValueError("Server configuration error: AWS credentials not found")
    except ClientError as e:
        error_code = e.response['Error']['Code']
        logger.error(f"AWS S3 error: {str(e)}")

        if error_code == 'AccessDenied':
            raise ValueError(
                "Profile picture upload is temporarily unavailable. Please try again later or contact support.")
        else:
            raise ValueError(f"Failed to upload file to S3: {str(e)}")
    except Exception as e:
        logger.error(f"Unexpected error during S3 upload: {str(e)}")
        raise ValueError(f"Failed to upload profile picture: {str(e)}")


if __name__ == "__main__":
    candidate = Candidate.objects.order_by("?").first()
    # print(extract_from_resume(candidate))
    query = natural_language_to_query("New graduate Python developer")
    print(query)
    rank = rank_candidates(query, Candidate.objects.all())
    for candidate, match in rank:
        print(candidate.experience_months, candidate.title,
              candidate.ai_highlights[2], match)
